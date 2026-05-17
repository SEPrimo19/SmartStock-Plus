#!/usr/bin/env python3
"""Generate SmartStock+ MCO 1 Presentation Scripts Word Document."""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

GREEN = RGBColor(0x1B, 0x5E, 0x20)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
BLUE = RGBColor(0x15, 0x65, 0xC0)


def add_title(text, size=24, color=GREEN):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = True
    return p


def add_subtitle(text, size=13, color=GRAY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_slide_header(slide_num, title, speaker, role):
    # Horizontal rule
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("─" * 75)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)

    # Slide number + title
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"SLIDE {slide_num}: {title}")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = GREEN

    # Speaker info
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Speaker: ")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = DARK
    run = p.add_run(speaker)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK

    # Role
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"Role: {role}")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = GRAY


def add_script(text):
    """Add script text as an indented, quoted paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    return p


def add_tip(text):
    """Add a tip/note in blue italic."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"Tip: {text}")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = BLUE


def add_dual_speaker(speaker, text):
    """For demo slide with two speakers."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(f"{speaker}: ")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = GREEN
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()  # spacer
doc.add_paragraph()
add_title("SmartStock+", size=28)
add_subtitle("Offline-First Inventory & Asset Management System", size=14, color=DARK)
doc.add_paragraph()
add_title("Presentation Scripts", size=22, color=DARK)
add_subtitle("MCO 1 — Mobile Programming 2 Midterm", size=13, color=GRAY)
doc.add_paragraph()
add_subtitle("18 Slides  |  12 Speakers  |  20–25 Minutes", size=12, color=GRAY)
doc.add_paragraph()

# Member list
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run("Team 2 Members")
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = GREEN

members = [
    "Jhon Clarence B. Rulona — Project Leader, Back-End Developer",
    "Cris Gerard O. Carpon — Project Leader, Back-End Developer",
    "Jasper Keith P. Teorica — Project Leader",
    "Mark Jacob S. Allero — UI/UX Designer",
    "John Dennis Y. Chan — UI/UX Designer",
    "Leopoldo C. Villaflores III — UI/UX Designer",
    "Dustin Aaron D. Tocayon — Back-End Developer, QA/Tester",
    "Aleah A. Cantiga — Quality Assurance/Tester",
    "Angel Mae S. Dura — Quality Assurance/Tester",
    "Heizel D. Panugaling — Quality Assurance/Tester",
    "Justin Jamaica P. Julaton — Quality Assurance/Tester",
    "Roselyn B. Deguino — Document Specialist",
    "Lou Ariane Mae B. Delos Reyes — Document Specialist",
    "Rynel Jay V. Vallejos — Document Specialist",
]
for m in members:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(m)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SPEAKER ASSIGNMENT TABLE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run("Speaker Assignment Summary")
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = GREEN

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run("12 members present (Roselyn Deguino and Aleah Cantiga are unavailable). All present members support Q&A.")
run.font.size = Pt(11)
run.font.color.rgb = GRAY

assignments = [
    ("1", "Title Slide", "Jhon Clarence B. Rulona"),
    ("2", "Project Overview", "Jasper Keith P. Teorica"),
    ("3", "Project Goal", "Rynel Jay V. Vallejos"),
    ("4", "Midterm Scope", "Cris Gerard O. Carpon"),
    ("5", "App Architecture Overview", "Leopoldo C. Villaflores III"),
    ("6", "Why Use Hilt?", "Dustin Aaron D. Tocayon"),
    ("7", "Hilt Module", "Lou Ariane Mae B. Delos Reyes"),
    ("8", "Data Model (Entity)", "Jhon Clarence B. Rulona"),
    ("9", "DAO (Data Access Object)", "Heizel D. Panugaling"),
    ("10", "Repository Pattern", "Cris Gerard O. Carpon"),
    ("11", "ViewModel", "Justin Jamaica P. Julaton"),
    ("12", "UI (Compose Screenshots)", "Mark Jacob S. Allero"),
    ("13", "Data Flow with Hilt", "John Dennis Y. Chan"),
    ("14", "Demonstration", "Jhon Clarence + Mark Jacob"),
    ("15", "Key Achievements", "Heizel D. Panugaling"),
    ("16", "Challenges Encountered", "Angel Mae S. Dura"),
    ("17", "Preparation for MCO 2", "Rynel Jay V. Vallejos"),
    ("18", "Conclusion", "Jasper Keith P. Teorica"),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for i, header in enumerate(["Slide", "Title", "Speaker"]):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

# Data rows
for slide_num, title, speaker in assignments:
    row = table.add_row()
    row.cells[0].text = slide_num
    row.cells[1].text = title
    row.cells[2].text = speaker
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(0.6)
    row.cells[1].width = Inches(2.8)
    row.cells[2].width = Inches(3.2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION HEADER
# ══════════════════════════════════════════════════════════════════════════════
add_title("Individual Presentation Scripts", size=20, color=GREEN)
add_subtitle("Read naturally — do not memorize word-for-word.\n"
             "Understand the key points so you can explain in your own words.",
             size=11, color=GRAY)

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 1
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(1, "Title Slide", "Jhon Clarence B. Rulona",
                 "Project Leader, Back-End Developer")

add_script(
    "Good day, everyone. We are Team 2, and today we present SmartStock+ — "
    "an Offline-First Inventory and Asset Management System."
)
add_script(
    "This is our midterm project for Mobile Programming 2. "
    "SmartStock+ is built with Kotlin, Jetpack Compose, Material Design 3, "
    "Hilt for dependency injection, and Room for local storage."
)
add_script(
    "Our team has 14 members — Project Leaders, Back-End Developers, "
    "UI/UX Designers, QA Testers, and Document Specialists. "
    "Everyone will present today."
)
add_script(
    "Today we will cover our app architecture, Room database, "
    "Hilt implementation, all nine screens, and a live demo."
)
add_script(
    "Now I will pass it to Jasper for our project overview."
)
add_tip("Keep it warm and confident. Make eye contact. This sets the tone.")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 2
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(2, "Project Overview", "Jasper Keith P. Teorica",
                 "Project Leader")

add_script(
    "SmartStock+ is an Android inventory and asset management app. "
    "It lets organizations track equipment, tools, and supplies on a mobile device."
)
add_script(
    "The key feature is offline-first design. "
    "Users can manage their entire inventory without internet. "
    "All data is stored locally using Room database."
)
add_script(
    "Our target users are school labs, IT departments, "
    "student organizations, small offices, and asset custodians."
)
add_script(
    "The problem is simple — paper logs and spreadsheets are error-prone, "
    "hard to maintain, and inaccessible during outages. "
    "SmartStock+ solves this with a reliable mobile solution."
)
add_script(
    "Rynel Jay will now discuss our project goals."
)

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 3
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(3, "Project Goal", "Rynel Jay V. Vallejos",
                 "Document Specialist")

add_script(
    "Our goals for SmartStock+ focus on building a solid foundation."
)
add_script(
    "First — build a scalable app with clean MVVM architecture "
    "and the Repository pattern to separate UI, logic, and data."
)
add_script(
    "Second — implement offline-first readiness using Room database, "
    "so the app works fully without internet."
)
add_script(
    "Third — use Hilt for dependency injection. "
    "This is Google's recommended DI framework for Android."
)
add_script(
    "Fourth — deliver role-based access control with Admin and Staff roles, "
    "each with different permissions."
)
add_script(
    "Fifth — prepare the architecture for cloud sync, "
    "which we will complete in MCO 2."
)
add_script(
    "Cris will now explain our midterm scope."
)
add_tip("Point to each goal on the slide as you mention it.")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 4
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(4, "Midterm Scope", "Cris Gerard O. Carpon",
                 "Project Leader, Back-End Developer")

add_script(
    "Let me clarify what MCO 1 includes and what is planned for MCO 2."
)
add_script(
    "In MCO 1 we have: "
    "a Room database at version 7 with six entities and a full migration chain. "
    "Hilt dependency injection using KSP. "
    "Nine Jetpack Compose screens with Material Design 3. "
    "MVVM with the Repository pattern."
)
add_script(
    "We also have authentication with Admin and Staff roles. "
    "Item usage tracking — users specify a destination when checking out, "
    "and provide a reason when returning. "
    "Image capture using camera or gallery. "
    "And barcode slash QR scanning."
)
add_script(
    "What is NOT in MCO 1: "
    "Supabase cloud integration, remote data sync, "
    "and usage reporting with date filters. "
    "These are planned for MCO 2."
)
add_script(
    "Leopoldo will now walk us through our architecture."
)

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 5
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(5, "App Architecture Overview", "Leopoldo C. Villaflores III",
                 "UI/UX Designer")

add_script(
    "As you can see in this diagram, our app has four layers, "
    "and Hilt connects them all."
)
add_script(
    "At the top — the UI Layer. "
    "Built entirely with Jetpack Compose and Material Design 3. "
    "It has nine screens. "
    "The UI observes state using collectAsStateWithLifecycle."
)
add_script(
    "Below that — the ViewModel Layer. "
    "We have InventoryViewModel and DashboardViewModel, "
    "both annotated with @HiltViewModel. "
    "This layer handles business logic, validation, and role checks."
)
add_script(
    "Next — the Repository Layer. "
    "InventoryRepository receives both the DAO and a CloudSyncDataSource. "
    "Right now the sync source is a NoOp stub. "
    "In MCO 2, we just swap it for Supabase."
)
add_script(
    "At the bottom — the Data Layer. "
    "Room database version 7, with six entities, "
    "35 DAO operations, and six migrations. "
    "All reads return Flow for reactive updates."
)
add_script(
    "On the right side — Hilt DI wires everything together at compile time. "
    "No manual factories needed."
)
add_script(
    "Dustin will now explain why we chose Hilt."
)
add_tip("Point to each layer on the diagram as you explain it.")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 6
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(6, "Why Use Hilt?", "Dustin Aaron D. Tocayon",
                 "Back-End Developer, QA/Tester")

add_script(
    "Hilt is Google's recommended dependency injection library for Android."
)
add_script(
    "First — compile-time safety. "
    "Hilt catches dependency errors during the build, not at runtime."
)
add_script(
    "Second — it is annotation-driven. "
    "@HiltViewModel, @Inject, and @Module replace hundreds of lines "
    "of manual factory code."
)
add_script(
    "Third — scoped lifecycles. "
    "Our database is a singleton for the whole app. "
    "ViewModels are scoped to their screens."
)
add_script(
    "Fourth — reduced boilerplate. "
    "No custom ViewModelProvider Factory classes needed."
)
add_script(
    "Fifth — scalability. "
    "In MCO 2, adding cloud sync and WorkManager is easy "
    "because Hilt handles new dependencies cleanly."
)
add_script(
    "One challenge: the Hilt Gradle plugin was incompatible with AGP 9.0.1. "
    "We resolved this with a KSP-only approach."
)
add_script(
    "Lou Ariane Mae will now show the actual Hilt module code."
)
add_tip("Point to each benefit card on the slide as you mention it.")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 7
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(7, "Hilt Module (Providing Dependencies)",
                 "Lou Ariane Mae B. Delos Reyes", "Document Specialist")

add_script(
    "This is our AppModule — the Hilt module that provides all dependencies."
)
add_script(
    "@Module and @InstallIn with SingletonComponent means "
    "these are app-wide singletons."
)
add_script(
    "The first function provides the Room database. "
    "It uses a factory pattern — AppDatabase.getDatabase — "
    "which builds the database with all six migrations inside."
)
add_script(
    "The second provides the InventoryDao from the database."
)
add_script(
    "The third provides a CloudSyncDataSource. "
    "For now it is a NoOp stub — a placeholder. "
    "In MCO 2, we replace this with the actual Supabase implementation."
)
add_script(
    "The fourth provides the Repository with both the DAO "
    "and the cloud sync source."
)
add_script(
    "The fifth provides a ConnectivityObserver "
    "that monitors network state for future sync."
)
add_script(
    "With this single module, Hilt resolves the entire dependency chain "
    "at compile time. "
    "Jhon Clarence will now walk through our data model."
)

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 8
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(8, "Data Model (Entity)", "Jhon Clarence B. Rulona",
                 "Project Leader, Back-End Developer")

add_script(
    "Our core entity is InventoryItem. "
    "It uses the @Entity annotation with a unique index on assetCode "
    "for barcode scanning."
)
add_script(
    "The fields cover everything for inventory management: "
    "asset code, name, description, category, quantity, "
    "in-use quantity, condition, status, location, "
    "image URI for photos, and timestamps."
)
add_script(
    "In total, we have six entities."
)
add_script(
    "InventoryItem is the core table. "
    "ItemHistory provides an audit trail with cascade delete. "
    "ItemUsageRecord tracks checkout details — "
    "destination, who checked it out, return reason, and status."
)
add_script(
    "LocalUser manages Admin and Staff accounts with unique email. "
    "CategoryEntity holds categories — seeded defaults plus custom ones. "
    "AssetStatusEntity stores Available, In-Use, Damaged, and Retired."
)
add_script(
    "Heizel will now continue with our DAO."
)
add_tip("Point to the entity cards on the right side of the slide.")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 9
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(9, "DAO (Data Access Object)", "Heizel D. Panugaling",
                 "Quality Assurance/Tester")

add_script(
    "Our InventoryDao has 35 operations in total."
)
add_script(
    "For reads — we get all items, search by name or asset code, "
    "filter by category or status, get history by item, "
    "get active usage records, and fetch users."
)
add_script(
    "Every read returns a Kotlin Flow. "
    "This means the UI updates automatically when data changes. "
    "No manual refresh needed."
)
add_script(
    "For writes — insert, update, and delete for items, history, "
    "usage records, users, categories, and statuses. "
    "These are all suspend functions, so they run on background threads."
)
add_script(
    "A key design: our search uses LIKE with wildcards on both "
    "name and asset code, so barcode scans find matches instantly."
)
add_script(
    "Cris will now explain the Repository pattern."
)

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 10
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(10, "Repository Pattern", "Cris Gerard O. Carpon",
                 "Project Leader, Back-End Developer")

add_script(
    "Our InventoryRepository is the single source of truth "
    "between UI and data."
)
add_script(
    "It takes two parameters: the InventoryDao and a CloudSyncDataSource. "
    "Both are provided by Hilt through our AppModule."
)
add_script(
    "The CloudSyncDataSource defaults to a NoOp implementation. "
    "This means cloud sync is stubbed for MCO 1. "
    "Every write operation — insert and update — "
    "queues the change to the sync source automatically."
)
add_script(
    "The key benefit: ViewModels never touch the DAO directly. "
    "They only talk to the Repository. "
    "So when we add Supabase in MCO 2, "
    "we only change the Repository layer. "
    "ViewModels and UI stay untouched."
)
add_script(
    "Justin Jamaica will now discuss our ViewModel."
)
add_tip("Point to the MCO 2 callout banner at the bottom of the slide.")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 11
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(11, "ViewModel", "Justin Jamaica P. Julaton",
                 "Quality Assurance/Tester")

add_script(
    "Our InventoryViewModel uses @HiltViewModel and @Inject constructor. "
    "Hilt creates and manages this ViewModel automatically."
)
add_script(
    "For state, we use StateFlow — not LiveData. "
    "Room Flows are converted with stateIn, "
    "using WhileSubscribed set to 5 seconds. "
    "This stops collection when no screen is listening, saving resources."
)
add_script(
    "For filtered results, we use Kotlin's combine operator "
    "to merge search, category filter, status filter, and sort setting "
    "into one reactive stream."
)
add_script(
    "The ViewModel also handles business logic. "
    "For example, useItem checks the user's role permissions, "
    "validates the quantity, updates in-use count, "
    "creates a usage record with the destination, and logs to history."
)
add_script(
    "returnItem does the reverse — closes usage records, "
    "records the return reason, and restores availability."
)
add_script(
    "All permission checks happen here, so the UI just calls the function."
)
add_script(
    "Mark Jacob will now show our app screens."
)

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 12
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(12, "UI (Compose Screenshots)", "Mark Jacob S. Allero",
                 "UI/UX Designer")

add_script(
    "Here are all nine screens, built with Jetpack Compose "
    "and Material Design 3."
)
add_script(
    "Starting top-left — the Splash Screen with our logo. "
    "Then the Login Screen — notice the bottom navigation is hidden here. "
    "Users must log in first."
)
add_script(
    "The Dashboard shows total, available, and in-use counts. "
    "The Inventory List has search, filters, sort, "
    "and swipe-to-delete for Admins."
)
add_script(
    "The Add/Edit form includes image capture from camera or gallery, "
    "and a plus icon to add custom categories."
)
add_script(
    "Item Detail shows the photo, full info, "
    "and action buttons: Use, Return, and Condition. "
    "Use asks for a destination location. "
    "Return asks for a reason. "
    "Condition offers New, Good, Fair, and Poor."
)
add_script(
    "We also have History for the audit trail, "
    "Camera Scan for barcodes, "
    "and Profile with account info, dark mode, and user management."
)
add_script(
    "John Dennis will now explain our data flow."
)
add_tip("Point to each screenshot as you describe it.")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 13
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(13, "Data Flow with Hilt", "John Dennis Y. Chan",
                 "UI/UX Designer")

add_script(
    "Let me walk through how data flows, "
    "using the 'Use Item' action as an example."
)
add_script(
    "Step one — the user taps Use and enters a quantity and destination."
)
add_script(
    "Step two — the ViewModel, injected by Hilt, "
    "checks if the user's role has permission. "
    "If yes, it validates the quantity."
)
add_script(
    "Step three — it calls the Repository, also injected by Hilt. "
    "The Repository updates the item in Room, "
    "inserts an ItemUsageRecord, and inserts an ItemHistory entry."
)
add_script(
    "Step four — Room persists the changes and emits an updated Flow."
)
add_script(
    "Step five — the StateFlow in the ViewModel picks up the change. "
    "The Compose UI collects it and recomposes automatically. "
    "The screen shows updated quantities instantly."
)
add_script(
    "The whole chain is reactive and automatic. "
    "Hilt wired it all at compile time."
)
add_script(
    "Now Jhon Clarence and Mark Jacob will give a live demo."
)
add_tip("Trace the arrow path on the diagram as you explain each step.")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 14
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(14, "Demonstration",
                 "Jhon Clarence B. Rulona + Mark Jacob S. Allero",
                 "Project Leader + UI/UX Designer")

add_dual_speaker("Jhon Clarence",
    "Let us give you a live demo of SmartStock+. Launching the app now."
)
add_dual_speaker("Mark Jacob",
    "Here is the Splash Screen. "
    "It transitions to Login — the bottom navigation is hidden. "
    "Users must authenticate first."
)
add_dual_speaker("Jhon Clarence",
    "Logging in as Admin. "
    "Now we are on the Dashboard with summary cards. "
    "Let me go to the Inventory List. "
    "Watch the search — results update in real time as I type. "
    "I can also filter by category and status."
)
add_dual_speaker("Mark Jacob",
    "Let me add a new item. "
    "I will select a category — and tap the plus icon to add a custom one. "
    "Now I will take a photo with the camera. "
    "Set quantity and condition, and save."
)
add_dual_speaker("Jhon Clarence",
    "The item appears in the list. "
    "Tapping it shows the detail with the photo. "
    "Now I tap Use — it asks for a destination location. "
    "I enter 'Faculty Room' and check out 2 units. "
    "The quantities update instantly."
)
add_dual_speaker("Mark Jacob",
    "Now returning one unit. "
    "The return dialog asks for a reason — 'Task completed.' "
    "Available quantity goes back up. "
    "I can also update condition — New, Good, Fair, or Poor."
)
add_dual_speaker("Jhon Clarence",
    "Now I log out and log in as Staff. "
    "Staff cannot add or delete items, but can use, return, and update condition. "
    "And all of this works completely offline — "
    "the entire app runs on the local Room database. "
    "Heizel will now present our key achievements."
)
add_tip("Rehearse the demo at least twice. Have backup screenshots ready.")

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 15
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(15, "Key Achievements", "Heizel D. Panugaling",
                 "Quality Assurance/Tester")

add_script(
    "As part of QA, our team verified every feature. "
    "Here are our key achievements."
)
add_script(
    "One — Room database version 7 with six entities "
    "and a complete migration chain. All migrations preserve data."
)
add_script(
    "Two — clean MVVM with Hilt DI using KSP. "
    "We resolved the AGP 9.0.1 compatibility issue."
)
add_script(
    "Three — fully offline-first. Every feature works without internet."
)
add_script(
    "Four — role-based authentication. "
    "Admin has full CRUD. Staff has limited permissions."
)
add_script(
    "Five — complete usage lifecycle. "
    "Checkout with destination, return with reason."
)
add_script(
    "Six — image capture with camera and gallery using Coil. "
    "Seven — barcode and QR scanning. "
    "Eight — adaptive UI with bottom nav and navigation rail. "
    "Nine — dark mode with Material Design 3. "
    "Ten — nine Compose screens with reactive StateFlow updates."
)
add_script(
    "Every MCO 1 requirement has been met and exceeded. "
    "Angel Mae will discuss the challenges."
)

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 16
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(16, "Challenges Encountered", "Angel Mae S. Dura",
                 "Quality Assurance/Tester")

add_script(
    "We encountered several challenges during development."
)
add_script(
    "First — Hilt and AGP 9.0.1 incompatibility. "
    "The Hilt Gradle plugin dropped support for BaseExtension API. "
    "We fixed this with a KSP-only approach — "
    "applying KSP directly without the Hilt plugin."
)
add_script(
    "Second — Room migration complexity. "
    "We went through seven schema versions. "
    "Each migration required careful ALTER TABLE statements "
    "to preserve existing data."
)
add_script(
    "Third — Flow inside coroutines. "
    "Using collect on a Room Flow inside viewModelScope.launch "
    "would suspend indefinitely. "
    "We fixed this by using .first() for one-shot queries."
)
add_script(
    "Fourth — combining multiple reactive streams. "
    "Search, filter, sort, and role state had to merge "
    "using Kotlin's combine operator."
)
add_script(
    "Fifth — camera integration required FileProvider setup "
    "and runtime permission handling across Android versions."
)
add_script(
    "Rynel Jay will now discuss our MCO 2 plans."
)

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 17
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(17, "Preparation for Final Term (MCO 2)",
                 "Rynel Jay V. Vallejos", "Document Specialist")

add_script(
    "For MCO 2, we have a clear roadmap."
)
add_script(
    "First — cloud sync using Supabase. "
    "Our Repository already abstracts the data layer. "
    "We just add a remote data source — "
    "the ViewModel layer will not change."
)
add_script(
    "Second — a dedicated usage tracking screen. "
    "Staff needs to see all checked-out items, their locations, "
    "and who checked them out. "
    "This will support monthly and annual date filters."
)
add_script(
    "Third — WorkManager for background sync. "
    "Pending changes will sync to the cloud even when the app is closed."
)
add_script(
    "Fourth — enhanced barcode scanning with auto-match "
    "and batch scanning."
)
add_script(
    "Fifth — UI polish with animations and empty state improvements."
)
add_script(
    "Jasper will now close our presentation."
)

# ══════════════════════════════════════════════════════════════════════════════
# SLIDE 18
# ══════════════════════════════════════════════════════════════════════════════
add_slide_header(18, "Conclusion", "Jasper Keith P. Teorica",
                 "Project Leader")

add_script(
    "That concludes our MCO 1 presentation for SmartStock+."
)
add_script(
    "To summarize — we delivered a fully functional offline-first "
    "inventory management app."
)
add_script(
    "Clean MVVM architecture with Hilt dependency injection. "
    "Room database version 7 with six entities, "
    "35 DAO operations, and a complete migration chain."
)
add_script(
    "Nine Jetpack Compose screens with Material Design 3. "
    "Role-based authentication with granular permissions. "
    "Item usage tracking with destinations and return reasons. "
    "Image capture and barcode scanning."
)
add_script(
    "Our architecture is designed for scalability — "
    "the Repository pattern means we plug in cloud sync for MCO 2 "
    "without restructuring the app."
)
add_script(
    "We are confident in our progress and ready for the final term."
)
add_script(
    "We are now open for questions. Thank you."
)
add_tip("Smile. Make eye contact. Pause before saying 'Thank you.'")

# ══════════════════════════════════════════════════════════════════════════════
# TIPS PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

add_title("Tips for All Speakers", size=18, color=GREEN)

tips = [
    ("Pace", "Each speaker has roughly 1–2 minutes. Do not rush. Pause between key points."),
    ("Eye Contact", "Face the audience, not the screen. Refer to the slide only to point at items."),
    ("Transitions", "Each speaker names the next speaker when handing off. This keeps it smooth."),
    ("Technical Terms", "When saying StateFlow, Room, DAO, Hilt, MVVM — briefly explain if the audience looks confused."),
    ("Key Numbers", "Remember these: 6 entities, 35 DAO operations, 6 migrations, 9 screens, 2 roles."),
    ("Demo Rehearsal", "Practice the live demo at least twice before the real presentation."),
    ("Demo Backup", "If the demo has technical issues, have screenshots ready as backup."),
    ("Q&A Prep", "Every member should prepare answers for questions about their slide and their role area."),
    ("Memorization", "Do not memorize word-for-word. Understand the key points and explain naturally."),
    ("Confidence", "You built this app. You know it. Speak with confidence."),
]

for i, (title, desc) in enumerate(tips):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{i+1}. {title}: ")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREEN
    run = p.add_run(desc)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK

# ══════════════════════════════════════════════════════════════════════════════
# Q&A CHEAT SHEET
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Q&A Quick Reference")
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = GREEN

qa_items = [
    ("Why offline-first?",
     "Target users like school labs often have unreliable internet. "
     "Room ensures the app always works."),
    ("Why Hilt over Koin or manual DI?",
     "Hilt has compile-time safety, is Google-recommended, "
     "and catches errors during build instead of runtime."),
    ("How does role-based access work?",
     "UserRole enum has 6 boolean permissions. "
     "ViewModel checks these before every operation."),
    ("What happens when internet comes back?",
     "CloudSyncDataSource queues changes. In MCO 2, WorkManager "
     "will sync pending items to Supabase automatically."),
    ("Why 6 migrations instead of destructive rebuild?",
     "Destructive migration deletes all user data. "
     "Our migrations preserve data across every version upgrade."),
    ("How does image capture work?",
     "FileProvider creates a temp URI. Camera or gallery writes to it. "
     "Coil loads and displays the image."),
    ("Why StateFlow instead of LiveData?",
     "StateFlow is coroutine-native, works better with Compose, "
     "and provides more control with operators like combine and map."),
]

for q, a in qa_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"Q: {q}")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"A: {a}")
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
output_path = r"C:\Users\Admin\Downloads\SmartStock+\SmartStock\Documentation\SmartStock_MCO1_Presentation_Scripts.docx"
doc.save(output_path)
print(f"Scripts document saved to: {output_path}")
